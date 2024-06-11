import os
from docx import Document
from docx.shared import Pt, Cm, Mm
import pandas as pd
import sqlite3


def conecBD():
    return sqlite3.connect('Sql_data/db_personas.db')

def gen_cons():
    conn = conecBD()
    consulta = """
    SELECT * 
    from personas
    INNER JOIN Salarios ON personas.id_rol = Salarios.id_salarios
    """
    DataFrame= pd.read_sql_query(consulta,conn)
    print(DataFrame)
    conn.close()

def example_contract(date: str, rol: str, address: str, rut: str, full_name: str, nationality: str, birth_date: str, profession: str, salary: str) -> str:
    document = Document()

    # Configuración de parámetros del documento
    font = document.styles['Normal'].font
    font.name = 'Book Antiqua'
    sections = document.sections
    for section in sections:
        section.page_height = Cm(35.56)
        section.page_width = Cm(21.59)
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)

    # Encabezado del documento (imagen header.png)
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture('imagenes/header.png')

    # Logo (imagen logo.png)
    document.add_picture('imagenes/logo.png')

    # Título del contrato
    h = document.add_paragraph('')
    h.add_run('CONTRATO DE PRESTACIÓN DE SERVICIOS A HONORARIOS\n').bold = True
    h.alignment = 1
    font.size = Pt(10)
    h.add_run('............................................................................................................................................................................................................................. ').bold = True

    # Contenido del contrato
    p = document.add_paragraph(f'En Temuco, a {str(date)}, entre la Corporación de Innovación y Desarrollo Tecnológico, Rut 78.898.766-4, representada por su Director General don(a) Roberto Gomez Bolainas, Cédula de Identidad Nº 10.678.990-2, ambos domiciliados en Caupolican 455 de esta ciudad, en adelante la “Corporación” y  {full_name}, de nacionalidad {nationality}, de profesión {profession}, nacido el {birth_date}, con domicilio en {address}, Cédula de Identidad N° {rut}, en adelante, el “Prestador de Servicios”, se ha convenido el siguiente contrato de prestación de servicios a honorarios: \n')
    font.size = Pt(8)
    p1 = document.add_paragraph('')
    p1.add_run('PRIMERO        :').bold = True
    p1.add_run('En el marco del acuerdo de servicios profesionales fechado el 11 de noviembre de 2020, establecido entre la Agencia Nacional de Estándares Educativos y la Corporación de Innovación y Desarrollo Tecnológico , y ratificado según la Resolución Exenta N°603 del 23 de noviembre de 2020, la Corporación encarga los servicios profesionales del Prestador de Servicios, para que ejecute la siguiente tarea en el proyecto "Evaluación de competencias específicas y metodologías de aprendizaje artificial 2020, ID 67703-20-JJ90."')
    p1.add_run('\n SEGUNDO        :').bold = True
    p1.add_run('El rol a desempeñar es de '+rol+'.')
    p1.add_run('\n TERCERO        :').bold = True
    p1.add_run('El plazo para la realización de la prestación de servicios encomendada será el '+str(date))
    p1.add_run('\n CUARTO        :').bold = True
    p1.add_run('Por el servicio profesional efectivamente realizado, se pagara un monto bruto variable, el cual corresponderá a cada rol dentro de la empresa capacitación, de acuerdo al siguiente detalle: ')
    table = document.add_table(rows=2, cols=2)
    table.alignment = 1
    hdr_cells0 = table.rows[0].cells
    hdr_cells0[0].text='Rol'
    hdr_cells0[1].text='Monto Bruto'
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = rol
    hdr_cells[1].text= salary
    p1.add_run('\n QUINTO        :').bold = True
    p1.add_run('El Prestador de Servicios acepta el encargo y las condiciones precedentes.')
    p1.add_run('\n SEXTO        :').bold = True
    p1.add_run(
        'El Prestador de Servicios está obligado a mantener la confidencialidad de todos los materiales utilizados, conforme al Acuerdo de Confidencialidad previamente establecido.')
    p1.add_run('\n En comprobante, previa lectura y ratificación, las partes firman.  ').bold = True

    # Firmas (modificado para incluir imagen de firma)
    table = document.add_table(rows=2, cols=2)
    table.alignment = 1
    hdr_cells0 = table.rows[0].cells[1].add_paragraph()
    r = hdr_cells0.add_run()
    r.add_picture('imagenes/firma.png')
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = '-----------------------------------------------------------\nEL PRESTADOR DE SERVICIOS'
    hdr_cells[1].text = '-----------------------------------------------------------\np. LA CORPORACION'

    # Pie de página (imagen footer1.png)
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run('Caupolican 0455, Temuco, Chile, www.corpoindet.cl\n')
    run.add_picture("imagenes/footer1.png")

    # Guardar el documento
    document.save(f'{full_name}.docx')

def BusquedaSing():
    conn = conecBD()
    while True:
        print('Desea realizar una busqueda por:')
        print('1. Rut')
        print('2. Nombre Completo')
        try:
            eleccion = int(input('Por favor ingrese el numero de su seleccion: '))

            if eleccion == 1:
                rut = str(input('Ingrese el rut de la persona: '))
                query = f"SELECT * FROM personas INNER JOIN Salarios ON personas.id_rol = Salarios.id_salarios WHERE personas.rut = ?"
                res = pd.read_sql_query(query, conn,params=(rut,))
                print(res)
                break
            elif eleccion == 2:
                nombre = input('Ingrese el Nombre completo de la persona: ').strip().lower()
                query = f"SELECT * FROM personas INNER JOIN Salarios ON personas.id_rol = Salarios.id_salarios WHERE personas.nombre_completo = ?"
                res = pd.read_sql_query(query, conn,params=(nombre,))
                print(res)
                break
        except ValueError:
            print('Ingrese un numero valido')
    conn.close()     

def BusquedaGrup():
    conn = conecBD()
    while True:
        print('Desea realizar una busqueda por: ')
        print('1.Sueldo' "\n" '2.Rol' "\n" '3.Nacionalidad' "\n" '4.Profesion')
        try:
            eleccion = int(input('Ingrese el numero de su eleccion: '))

            if eleccion == 1:
                sueldo = int(input('Ingrese el sueldo de la persona: '))
                query = f'SELECT * FROM personas INNER JOIN Salarios ON personas.id_rol = Salarios.id_salarios WHERE Salarios.Sueldo = ?'
                res = pd.read_sql_query(query, conn,params=(sueldo,))
                print(res)
                break
            elif eleccion == 2:
                rol = input('Ingresa el rol de la persona: ').strip().capitalize()
                query =f'SELECT * FROM personas INNER JOIN Salarios ON personas.id_rol = Salarios.id_salarios WHERE Salarios.Rol = ?'
                res = pd.read_sql_query(query,conn,params=(rol,))
                print(res)
                break
            elif eleccion == 3:
                nacionalidad = input('Ingrese la nacionalidad de la persona: ').strip().capitalize()
                query = f"SELECT * FROM personas INNER JOIN Salarios ON personas.id_rol = Salarios.id_salarios WHERE personas.Nacionalidad = ?"
                res = pd.read_sql_query(query,conn, params=(nacionalidad,))
                print(res)
                break
            elif eleccion == 4:
                profesion = input('Ingresa la profesion de la persona: ').strip().lower()
                query = f'SELECT * FROM personas INNER JOIN Salarios ON personas.id_rol = Salarios.id_salarios WHERE personas.profesion = ?'
                res = pd.read_sql_query(query,conn,params=(profesion,))
                print(res)
                break
            else:
                print('Opcion no valida vuelva a intentar ')
        except ValueError:
            print('Ingrese un numero valido')
    conn.close()

def menu():
    conn = conecBD()
    try:
        while True:
            print('Que tipo de busqueda desea ralizar: ')
            print('1. Busqueda singular (1 persona)')
            print('2. Busqueda por grupo (Varias personas)')
            try:
                filtro = int(input('Ingresa el numero de tu eleccion: '))

                if filtro == 1:
                    BusquedaSing()
                    break
                elif filtro == 2:
                    BusquedaGrup()
                    break
                else:
                 print('opcion no valida vuelva a intentar')
            except ValueError:
             print('Ingrese un numero valido')
    finally:
        conn.close()

while True:
    menu()
    repito = input('Desea realizar otra consulta Si/No: ').strip().lower()

    if repito != "si":
        print('Nos vemo')
        break